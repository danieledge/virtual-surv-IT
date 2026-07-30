"""
scripts/render_docx.py - render a Markdown artifact to a Word (.docx) document.

Opt-in third output format alongside the required .md + .html (docs/WAYS-OF-WORKING.md
"Document-control standard"): non-technical stakeholders redlining a BRD/FSD/delivery
report in Word can't comfortably do that from Markdown or a static HTML page. `.md` stays
the ONE authored source - this and render_html.py are both pure derivations from it, so a
".docx" is never hand-edited and never drifts from the artifact's real content.

Reuses render_html.markdown_to_safe_html() for the parse + XSS-sanitisation pass (one
Markdown parse, one allow-list, both output formats), then walks that same sanitised HTML
fragment with the stdlib html.parser into a python-docx Document - no second Markdown
parser to maintain.

python-docx is an OPTIONAL dependency (requirements-dev.txt), same tier as Markdown/bleach
for render_html.py: absent -> a clear ImportError-derived message, never a crash elsewhere.

Usage:
  python -m scripts.render_docx artifacts/<slug>/BRD-spoofing.md
  python -m scripts.render_docx artifacts/<slug>/BRD-spoofing.md --out BRD-spoofing.docx
"""

from __future__ import annotations

import argparse
import re
from html.parser import HTMLParser
from pathlib import Path
from typing import Optional

from scripts.render_html import _title_from, markdown_to_safe_html

try:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    _DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover - optional dependency
    _DOCX_AVAILABLE = False

_HEADING_TAGS = {"h1", "h2", "h3", "h4", "h5", "h6"}
_INLINE_STYLE_TAGS = {"strong", "b", "em", "i", "code", "s", "del", "sup", "sub"}
_BLOCK_TAGS = {"p", "li", "blockquote", "pre", "hr"}


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """python-docx has no public hyperlink API (as of 1.x) - this is the standard low-level
    recipe: a w:hyperlink element wrapping a styled run, wired to a relationship id."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = docx.oxml.OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run_el = docx.oxml.OxmlElement("w:r")
    rpr = docx.oxml.OxmlElement("w:rPr")
    color = docx.oxml.OxmlElement("w:color")
    color.set(qn("w:val"), "0969DA")
    rpr.append(color)
    underline = docx.oxml.OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run_el.append(rpr)
    text_el = docx.oxml.OxmlElement("w:t")
    text_el.text = text
    run_el.append(text_el)
    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)


class _DocxBuilder(HTMLParser):
    """Walks a sanitised HTML fragment (from markdown_to_safe_html) into a python-docx
    Document. Deliberately covers the same feature set render_html.py's CSS styles for -
    headings, paragraphs, inline emphasis, lists, tables, blockquotes, code, links, rules -
    not the full HTML spec: the input is always OUR OWN Markdown output, never arbitrary
    third-party HTML, so the tag set is small and known."""

    def __init__(self, document) -> None:
        super().__init__(convert_charrefs=True)
        self.doc = document
        self._stack: list[str] = []
        self._bold = 0
        self._italic = 0
        self._code = 0
        self._strike = 0
        self._para = None
        self._list_stack: list[str] = []  # "ul" | "ol" per nesting level
        self._table = None
        self._row_cells: list = []
        self._cell_para = None
        self._link_href: Optional[str] = None
        self._link_text_parts: list[str] = []
        self._pre_lines: list[str] = []
        self._in_pre = False

    # -- paragraph helpers -------------------------------------------------
    def _ensure_para(self, style: Optional[str] = None):
        if self._para is None:
            self._para = self.doc.add_paragraph(style=style)
        return self._para

    def _close_para(self) -> None:
        self._para = None

    def _add_run(self, text: str) -> None:
        if not text:
            return
        target = self._cell_para if self._cell_para is not None else self._ensure_para()
        run = target.add_run(text)
        run.bold = self._bold > 0
        run.italic = self._italic > 0
        run.strike = self._strike > 0
        if self._code > 0:
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)

    # -- HTMLParser overrides -----------------------------------------------
    def handle_starttag(self, tag: str, attrs) -> None:
        attrs_d = dict(attrs)
        self._stack.append(tag)
        if tag in _HEADING_TAGS:
            level = min(int(tag[1]), 4)  # docx built-in styles go Heading 1-9; cap at 4 here
            self._para = self.doc.add_heading(level=level)
        elif tag == "p":
            # Markdown always wraps blockquote text in a nested <p> - style THAT
            # paragraph as the quote rather than the (now-empty) blockquote tag itself.
            in_quote = "blockquote" in self._stack[:-1]
            self._para = self.doc.add_paragraph(style="Intense Quote" if in_quote else None)
        elif tag == "blockquote":
            pass  # styling is applied to the nested <p>, not this wrapper
        elif tag in ("strong", "b"):
            self._bold += 1
        elif tag in ("em", "i"):
            self._italic += 1
        elif tag in ("s", "del"):
            self._strike += 1
        elif tag == "code":
            self._code += 1
        elif tag == "pre":
            self._in_pre = True
            self._pre_lines = []
        elif tag in ("ul", "ol"):
            self._list_stack.append(tag)
        elif tag == "li":
            depth = max(len(self._list_stack) - 1, 0)
            kind = self._list_stack[-1] if self._list_stack else "ul"
            style = "List Bullet" if kind == "ul" else "List Number"
            if depth:
                style = f"{style} {min(depth + 1, 3)}"
            try:
                self._para = self.doc.add_paragraph(style=style)
            except KeyError:  # depth-styled variant not in the default template
                self._para = self.doc.add_paragraph(
                    style="List Bullet" if kind == "ul" else "List Number"
                )
        elif tag == "table":
            self._table = None  # column count known only once the first row is seen
            self._pending_rows: list[list[str]] = []
        elif tag == "tr":
            self._row_cells = []
        elif tag in ("th", "td"):
            self._row_cells.append("")
        elif tag == "a":
            self._link_href = attrs_d.get("href")
            self._link_text_parts = []
        elif tag == "hr":
            hr_para = self.doc.add_paragraph()
            hr_para.add_run("_" * 40)

    def handle_endtag(self, tag: str) -> None:
        if self._stack and self._stack[-1] == tag:
            self._stack.pop()
        if tag in _HEADING_TAGS or tag in ("p", "blockquote", "li"):
            self._close_para()
        elif tag in ("strong", "b"):
            self._bold = max(0, self._bold - 1)
        elif tag in ("em", "i"):
            self._italic = max(0, self._italic - 1)
        elif tag in ("s", "del"):
            self._strike = max(0, self._strike - 1)
        elif tag == "code":
            self._code = max(0, self._code - 1)
        elif tag == "pre":
            self._in_pre = False
            block = self.doc.add_paragraph()
            run = block.add_run("\n".join(self._pre_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            self._pre_lines = []
        elif tag in ("ul", "ol"):
            if self._list_stack:
                self._list_stack.pop()
        elif tag in ("th", "td"):
            pass  # row_cells already carries the accumulated text
        elif tag == "tr":
            self._pending_rows.append(list(self._row_cells))
            self._row_cells = []
        elif tag == "table":
            self._flush_table()
        elif tag == "a":
            text = "".join(self._link_text_parts).strip()
            if self._link_href and text:
                target = self._cell_para if self._cell_para is not None else self._ensure_para()
                _add_hyperlink(target, self._link_href, text)
            elif text:
                self._add_run(text)
            self._link_href = None
            self._link_text_parts = []

    def handle_data(self, data: str) -> None:
        if self._in_pre:
            self._pre_lines.append(data.rstrip("\n"))
            return
        if self._link_href is not None:
            self._link_text_parts.append(data)
            return
        if self._row_cells:
            self._row_cells[-1] += data
            return
        # Whitespace-only text BETWEEN block elements (markdown's own pretty-printed
        # indentation, e.g. the newline between </h1> and <blockquote>) is not content -
        # skip it rather than materialising an empty paragraph. Whitespace inside an
        # already-open paragraph/heading (self._para is not None) is real inter-word
        # spacing and must be kept.
        if not data.strip() and self._para is None:
            return
        self._add_run(data)

    def _flush_table(self) -> None:
        rows = getattr(self, "_pending_rows", [])
        if not rows:
            return
        n_cols = max(len(r) for r in rows)
        table = self.doc.add_table(rows=0, cols=n_cols)
        table.style = "Light Grid Accent 1"
        for r_idx, row in enumerate(rows):
            cells = table.add_row().cells
            for c_idx in range(n_cols):
                text = row[c_idx] if c_idx < len(row) else ""
                cells[c_idx].text = text.strip()
                if r_idx == 0:
                    for p in cells[c_idx].paragraphs:
                        for run in p.runs:
                            run.bold = True
        self._pending_rows = []


_TABLE_ROW_RE = re.compile(r"<tr\b")  # unused placeholder kept out; table state lives in handler


def render_docx(md_text: str, title: str, source: str = "") -> "docx.Document":
    if not _DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx is required for .docx export: pip install -r requirements-dev.txt"
        )
    safe_body = markdown_to_safe_html(md_text, html_links=False)
    document = docx.Document()
    document.core_properties.title = title
    if source:
        document.core_properties.comments = f"Rendered from {source}"

    header = document.sections[0].header
    header_p = header.paragraphs[0]
    header_p.text = "Compliance Surveillance Engineering"
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    builder = _DocxBuilder(document)
    builder.feed(safe_body)

    footer = document.sections[0].footer
    footer_p = footer.paragraphs[0]
    footer_p.text = "Generated from Markdown by the compliance surveillance engineering team." + (
        f" Source: {source}" if source else ""
    )
    return document


def main() -> None:
    ap = argparse.ArgumentParser(description="Render a Markdown artifact to a Word (.docx) file.")
    ap.add_argument("src", type=Path, help="path to the .md artifact")
    ap.add_argument("--out", type=Path, help="output .docx path (default: alongside the .md)")
    args = ap.parse_args()

    md_text = args.src.read_text(encoding="utf-8")
    out = args.out or args.src.with_suffix(".docx")
    document = render_docx(md_text, _title_from(md_text, args.src.stem), source=args.src.name)
    document.save(str(out))
    print(f"Rendered {args.src} -> {out}")


if __name__ == "__main__":
    main()
